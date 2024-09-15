import cv2
import numpy as np
import mediapipe as mp
import keyboard
import sys
import win32api
from collections import deque
import pyautogui
import pygetwindow
from PIL import Image
from datetime import datetime
import os
import subprocess
import screenshot
import copyfunc
import aspose.words as aw
import docx
#doc = aw.Document()
#builder = aw.DocumentBuilder(doc)
#font = builder.font
#font.size = 32
#font.bold = True
#font.name = "Arial"
#builder.writeln("Subject Lecture Notes")
#font.size = 14
#font.bold = False
#font.name = "Arial"
#y= datetime.today().strftime('%Y-%m-%d')
#builder.writeln(f"Date:{y}")
#no_of_files= len(os.listdir("C:\\Users\\ameya\\OneDrive\\Desktop\\2023-01-26"))
#for i in range(1,no_of_files+1):
    #builder.insert_image(f"C:\\Users\\ameya\\OneDrive\\Desktop\\2023-01-26\\{i}.png")

#doc.save("C:\\Files\\sample_output.docx")
def wordnotes():
    doc=docx.Document()
    y= datetime.today().strftime('%Y-%m-%d')
    doc.add_heading('Subject Notes',0)
    doc.add_heading(f'Date: {y}')
    no_of_files= len(os.listdir('C:\\Users\\ameya\\OneDrive\\Desktop\\2023-04-03'))
    for i in range(1,no_of_files):
        doc.add_picture(f'C:/Users/ameya/OneDrive/Desktop/2023-04-03/{i}.png')
    doc.save('doc121.docx')
    print("flag test case 1 confirm ")











